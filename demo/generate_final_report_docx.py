"""Render the final report (docs/final_report/SABERCAST_FINAL_REPORT.md) as a Word doc.

Output: docs/final_report/Sabercast_Final_Report.docx

Re-uses the markdown→docx renderer from generate_build_log_docx. Inserts the
architecture PNG (docs/architecture_diagram.png) immediately after the § 3
"System architecture" heading.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "demo"))

from generate_build_log_docx import (   # noqa: E402
    _add_heading,
    _add_horizontal_rule,
    _add_inline_runs,
    _add_table,
    _is_table_row,
    _parse_table_rows,
)
import re


MD_PATH       = ROOT / "docs" / "final_report" / "SABERCAST_FINAL_REPORT.md"
OUT_PATH      = ROOT / "docs" / "final_report" / "Sabercast_Final_Report.docx"
ARCH_PNG      = ROOT / "docs" / "architecture_diagram.png"


def _embed_architecture(doc: Document) -> None:
    """Add the architecture PNG centered, sized to fit page width."""
    if not ARCH_PNG.exists():
        print(f"  WARN: {ARCH_PNG.name} not found; skipping embed")
        return
    p = doc.add_paragraph()
    p.alignment = 1   # WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(ARCH_PNG), width=Inches(6.4))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = 1
    r = cap.add_run("Figure 1. Sabercast architecture (Mermaid source in docs/architecture_diagram.md). "
                    "Every box labeled with the specific model or data store doing the work.")
    r.font.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


_IMG_LINE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def _embed_image(doc: Document, alt: str, src: str) -> None:
    """Embed an image referenced from a markdown line.

    Resolution order:
      1. relative to the markdown file's directory (matches GitHub semantics)
      2. relative to the project root
      3. relative to docs/
    """
    md_dir = MD_PATH.parent
    candidates = [
        (md_dir / src).resolve(),
        ROOT / src,
        ROOT / "docs" / src,
    ]
    img_path = next((p for p in candidates if p.exists()), None)
    if img_path is None:
        print(f"  WARN: image not found: {src}")
        print(f"        tried: {[str(c) for c in candidates]}")
        return
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(6.4))
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = 1
        r = cap.add_run(alt)
        r.font.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def render_md_with_arch_embed(md_text: str, doc: Document) -> None:
    """Markdown -> Word docx with:
       * embedded architecture PNG after the '## 3. System architecture' heading
       * inline markdown image syntax ![caption](path) handled at line start
    """
    lines = md_text.splitlines()
    i = 0
    arch_inserted = False

    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        # Inline image line: ![caption](path)
        m_img = _IMG_LINE.match(line)
        if m_img:
            _embed_image(doc, m_img.group(1).strip(), m_img.group(2).strip())
            i += 1
            continue

        if line.strip() in {"---", "***", "___"}:
            _add_horizontal_rule(doc)
            i += 1
            continue

        if line.startswith("# "):
            _add_heading(doc, line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            _add_heading(doc, line[3:].strip(), level=2)
            i += 1
            # Embed the architecture diagram right after § 3
            if (not arch_inserted and
                line.startswith("## 3.")):
                _embed_architecture(doc)
                arch_inserted = True
            continue
        if line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            _add_inline_runs(p, line[2:].strip())
            i += 1
            continue

        if _is_table_row(line):
            header, rows, i = _parse_table_rows(lines, i)
            _add_table(doc, header, rows)
            continue

        if line.lstrip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, line.lstrip())
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\s*\d+\.\s+", "", line))
            i += 1
            continue

        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            lines[i].startswith(("# ", "## ", "### ", "> ", "- ", "* ", "!["))
            or _is_table_row(lines[i])
            or re.match(r"^\s*\d+\.\s", lines[i])
            or lines[i].strip() in {"---", "***", "___"}
        ):
            para_lines.append(lines[i].rstrip())
            i += 1
        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(para_lines))


def main() -> None:
    if not MD_PATH.exists():
        raise SystemExit(f"Final report markdown not found at {MD_PATH}")
    md_text = MD_PATH.read_text(encoding="utf-8")

    doc = Document()
    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    render_md_with_arch_embed(md_text, doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"saved {OUT_PATH}")


if __name__ == "__main__":
    main()
