"""Render the append-only build log (docs/BUILD_LOG.md) as a Word document.

Output: docs/Sabercast_Build_Log.docx

Idempotent — re-run after appending a new entry to BUILD_LOG.md and the docx
regenerates with the new content. Past entries are preserved verbatim because
the markdown itself is append-only.

Markdown features handled:
  * Top-level title (single `#`)
  * Entry headers (`##`)
  * Bold paragraph leaders (`**Goal:**`, `**What was built**`, `**What it produced**`)
  * Bullet lists (`- ` and `* `)
  * Pipe-delimited markdown tables
  * Horizontal rules (`---`) become small page-section separators
  * Inline `code` is rendered as monospace
  * Bold (`**x**`) inline is preserved
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT       = Path(__file__).resolve().parent.parent
MD_PATH    = ROOT / "docs" / "BUILD_LOG.md"
OUT_PATH   = ROOT / "docs" / "Sabercast_Build_Log.docx"


# ── Style helpers ────────────────────────────────────────────────────────────
def _set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


_INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_INLINE_CODE = re.compile(r"`([^`]+)`")
_INLINE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _add_inline_runs(paragraph, text: str, base_size: int = 10, base_font: str = "Arial") -> None:
    """Walk a markdown-flavored line, emitting runs with bold / code / link styling.

    The tokenizer is intentionally simple — it splits on **bold**, `code`, and
    [link](url) one pass at a time. Nested or overlapping styles are not
    supported; they don't appear in the build log.
    """
    # Strip leading-bullet artifacts if present
    if text.startswith(("- ", "* ")):
        text = text[2:]

    pos = 0
    pattern = re.compile(
        r"(\*\*[^*]+\*\*)"      # bold
        r"|(`[^`]+`)"            # code
        r"|(\[[^\]]+\]\([^)]+\))" # link
    )
    for m in pattern.finditer(text):
        # Plain text segment before this match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.name = base_font
            run.font.size = Pt(base_size)
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 0.5)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        else:
            # link [text](url) — render text, ignore url for now (could add hyperlink)
            link_m = _INLINE_LINK.match(token)
            if link_m:
                run = paragraph.add_run(link_m.group(1))
                run.font.name = base_font
                run.font.size = Pt(base_size)
                run.font.color.rgb = RGBColor(0x1A, 0x55, 0xB7)
                run.font.underline = True
        pos = m.end()
    # Trailing text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = base_font
        run.font.size = Pt(base_size)


def _add_heading(doc, text: str, level: int) -> None:
    p = doc.add_paragraph()
    sizes = {1: 20, 2: 15, 3: 12}
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(sizes.get(level, 12))
    run.font.bold = True
    if level == 1:
        run.font.color.rgb = RGBColor(0x10, 0x2A, 0x55)
    else:
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def _add_table(doc, header: list[str], rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.autofit = False
    table.allow_autofit = False
    content_width = Inches(6.4)
    col_w = Inches(6.4 / len(header))

    # Header row
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.width = col_w
        cell.text = ""
        r = cell.paragraphs[0].add_run(h.strip())
        r.font.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        _set_cell_shading(cell, "EEEEEE")

    # Data rows
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= len(header):
                continue
            cell = table.rows[r_idx].cells[c_idx]
            cell.width = col_w
            cell.text = ""
            _add_inline_runs(cell.paragraphs[0], cell_text.strip(), base_size=9.5)


def _add_horizontal_rule(doc) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:color"), "B0B0B0")
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Markdown parser (single-pass, line-by-line) ──────────────────────────────
def _is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and "|" in line.strip()[1:]


def _parse_table_rows(lines: list[str], i: int) -> tuple[list[str], list[list[str]], int]:
    """Consume a sequence of table lines starting at index i; return
    (header, data_rows, next_index)."""
    header_line = lines[i].strip()
    separator   = lines[i + 1].strip() if i + 1 < len(lines) else ""
    header = [c.strip() for c in header_line.strip("|").split("|")]
    rows: list[list[str]] = []
    j = i + 2
    if not (separator.startswith("|") and set(separator.replace("|", "").strip()) <= set("-: ")):
        # Not a real table — abort
        return header, rows, i + 1
    while j < len(lines) and _is_table_row(lines[j]):
        cells = [c.strip() for c in lines[j].strip("|").split("|")]
        rows.append(cells)
        j += 1
    return header, rows, j


def render_md_to_docx(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    in_list = False
    while i < len(lines):
        line = lines[i].rstrip()

        # Blank line → close any open paragraph context
        if not line.strip():
            in_list = False
            i += 1
            continue

        # Horizontal rule
        if line.strip() in {"---", "***", "___"}:
            _add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            _add_heading(doc, line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            _add_heading(doc, line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            _add_inline_runs(p, line[2:].strip())
            i += 1
            continue

        # Table
        if _is_table_row(line):
            header, rows, i = _parse_table_rows(lines, i)
            _add_table(doc, header, rows)
            continue

        # List item
        if line.lstrip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, line.lstrip())
            i += 1
            continue

        # Sub-numbered list (e.g. "  1. " or "  2. ")
        if re.match(r"^\s*\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\s*\d+\.\s+", "", line))
            i += 1
            continue

        # Regular paragraph — collect consecutive non-blank, non-special lines
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            lines[i].startswith(("# ", "## ", "### ", "> ", "- ", "* "))
            or _is_table_row(lines[i])
            or re.match(r"^\s*\d+\.\s", lines[i])
            or lines[i].strip() in {"---", "***", "___"}
        ):
            para_lines.append(lines[i].rstrip())
            i += 1
        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(para_lines))


def main() -> None:
    if not MD_PATH.exists():
        raise SystemExit(f"BUILD_LOG.md not found at {MD_PATH}")
    md_text = MD_PATH.read_text(encoding="utf-8")

    doc = Document()
    # US Letter, 1in margins
    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    render_md_to_docx(md_text, doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"saved {OUT_PATH}")


if __name__ == "__main__":
    main()
